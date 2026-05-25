"""
Génération, fusion et traduction du PV.
Compatible Ollama (modèles locaux) ET API Claude Anthropic.
"""
import json
import os
from datetime import datetime
from pathlib import Path
from typing import Optional
import re

import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

today = date.today().strftime("%d/%m/%Y")

SYSTEM_PV = """Tu es un expert en rédaction de procès-verbaux institutionnels pour des comités bancaires.
Tu rédiges en français administratif formel. Tu ne génères QUE du JSON valide, sans explication."""


# ─── Template PV ───────────────────────────────────────────────────────────────

EMPTY_PV_TEMPLATE = {
    "numero": "PV-2025-001",
    "titre": "Procès-Verbal de Réunion",
    "type_reunion": "Comité de Direction",
    "date": "",
    "heure_debut": "",
    "heure_fin": "",
    "lieu": "",
    "president_seance": "",
    "redacteur": "",
    "participants": [],
    "excuses": [],
    "ordre_du_jour": [],
    "points": [],
    "decisions": [],
    "plan_action": [],
    "prochaine_reunion": {
        "date": "",
        "lieu": "",
        "points_previsionnels": []
    },
    "approbation": {
        "date": "",
        "signataire": "",
        "statut": "En attente"
    }
}

POINT_TEMPLATE = {
    "numero": "1",
    "titre": "",
    "expose": "",
    "discussion": "",
    "remarques": [],
    "conclusion": ""
}

ACTION_TEMPLATE = {
    "id": "A01",
    "action": "",
    "responsable": "",
    "echeance": "",
    "statut": "En cours",
    "priorite": "Normale"
}

# ─── Config backend IA ─────────────────────────────────────────────────────────

today = datetime.now().strftime("%d/%m/%Y")
OLLAMA_URL        = os.environ.get("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL      = os.environ.get("OLLAMA_MODEL", "qwen2.5:7b")
EXPORT_DIR = Path(__file__).parent.parent / "exports"
EXPORT_DIR.mkdir(exist_ok=True)


# ─── Extraction JSON robuste ───────────────────────────────────────────────────

def _extract_json(text: str) -> dict | None:
    """
    Tente d'extraire un objet JSON depuis n'importe quelle réponse LLM.
    Gère : JSON pur, ```json ... ```, JSON noyé dans du texte, réponse partielle.
    """
    if not text or not text.strip():
        return None

    text = text.strip()

    # 1. Tentative directe
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 2. Extraire bloc ```json ... ```
    md_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
    if md_match:
        try:
            return json.loads(md_match.group(1))
        except json.JSONDecodeError:
            pass

    # 3. Extraire le premier { ... } de la réponse (même partiel dans du texte)
    brace_match = re.search(r'\{.*\}', text, re.DOTALL)
    if brace_match:
        try:
            return json.loads(brace_match.group(0))
        except json.JSONDecodeError:
            pass

    # 4. Réparer le JSON tronqué (fermer les accolades manquantes)
    if brace_match:
        candidate = brace_match.group(0)
        # Compter les accolades ouvertes/fermées
        opens  = candidate.count('{')
        closes = candidate.count('}')
        if opens > closes:
            candidate = candidate + '}' * (opens - closes)
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                pass

    return None


def _build_fallback_pv(extracted: dict) -> dict:
    """
    Fallback sans LLM : construit un PV structuré complet
    en mappant chaque point ODJ à ses données de slide.
    """
    slides = extracted.get("slides", [])
    odj    = _extract_odj(slides) or extracted.get("ordre_du_jour", [])

    # Participants
    participants_raw = extracted.get("participants", [])
    participants = [
        {"nom": p, "fonction": "", "present": True} if isinstance(p, str)
        else {"nom": p.get("nom", ""), "fonction": p.get("fonction", ""), "present": True}
        for p in participants_raw
    ]

    # Titre
    titre_reunion = ""
    for s in slides:
        t = s.get("titre") or ""
        if t and not t.isdigit() and len(t) > 5 and "ordre" not in t.lower():
            titre_reunion = t
            break
    titre_reunion = titre_reunion or extracted.get("meta", {}).get("titre_reunion", "Comité des Risques")

    # Participants pour introduction
    participants_intro = "\n".join(
        f"• {p['nom']}{' : ' + p['fonction'] if p['fonction'] else ''}"
        for p in participants
    ) or "• [À compléter]"

    # Introduction
    odj_liste = "\n".join(f"{i+1}. {t}" for i, t in enumerate(odj))
    introduction = (
        f"Le Comité des risques de Wifak Banque s'est réuni le {today} à 09h00 "
        f"au Centre d'Affaires de Tunis, sur invitation de son Président, "
        f"pour délibérer sur l'ordre du jour suivant :\n\n{odj_liste}\n\n"
        f"Étaient présents :\n{participants_intro}\n\n"
        f"La feuille de présence a été établie et dûment signée par les membres présents.\n\n"
        f"Après avoir constaté la présence requise des membres, le Président a ouvert "
        f"les travaux du Comité en souhaitant la bienvenue aux présents."
    )

    # Génération des points ODJ
    points = []
    for i, titre_odj in enumerate(odj):
        num = i + 1

        if num == 1:
            # Point 1 toujours identique
            points.append({
                "numero":     "1",
                "titre":      titre_odj,
                "expose":     "",
                "discussion": (
                    "Les membres présents ont approuvé l'ordre du jour et ont procédé "
                    "à l'examen des points inscrits sur lesquels ils ont délibéré."
                ),
                "conclusion": "",
            })
            continue

        # Données brutes de la slide correspondante
        data_bloc = _build_point_data(num, titre_odj, slides)

        # Expose générique + données
        expose = (
            f"Dans le cadre du point relatif à « {titre_odj} », "
            f"le Comité a procédé à l'examen des indicateurs et données présentés."
        )

        points.append({
            "numero":     str(num),
            "titre":      titre_odj,
            "expose":     expose,
            "discussion": data_bloc,
            "conclusion": (
                f"Le Comité a pris acte des informations relatives au point {num} "
                f"et a délibéré en conséquence."
            ),
        })

    # Plan d'action depuis les tableaux de la dernière slide de contenu
    plan_action = []
    action_idx = 1
    for slide in slides:
        for t in slide.get("tableaux", []):
            lignes = t.get("lignes", [])
            if not lignes:
                continue
            headers = [h.lower() for h in lignes[0]]
            # Détecte si c'est un tableau d'actions
            if any("action" in h for h in headers):
                col = {h: i for i, h in enumerate(headers)}
                for row in lignes[1:]:
                    if not any(row):
                        continue
                    plan_action.append({
                        "id":          f"A{str(action_idx).zfill(2)}",
                        "action":      row[col.get("action", 1)] if len(row) > col.get("action", 1) else "",
                        "responsable": row[col.get("responsable", 2)] if "responsable" in col and len(row) > col["responsable"] else "",
                        "echeance":    row[col.get("échéance", col.get("echeance", 3))] if len(row) > 3 else "",
                        "statut":      row[col.get("statut", 5)] if "statut" in col and len(row) > col.get("statut", 5) else "En cours",
                        "priorite":    row[col.get("priorité", col.get("priorite", 4))] if len(row) > 4 else "Normale",
                    })
                    action_idx += 1

    # Prochaine réunion
    prochaine_date = ""
    for slide in slides:
        for c in slide.get("contenu", []):
            import re
            m = re.search(r"(\w+\s+\d{4}|\d{1,2}/\d{1,2}/\d{4})", c)
            if m and ("prochain" in c.lower() or "juillet" in c.lower()):
                prochaine_date = m.group(1)
                break

    return {
        "numero":           f"PV-{today.replace('/', '-')}",
        "titre":            titre_reunion,
        "type_reunion":     "Comité des Risques",
        "date":             today,
        "heure_debut":      "09h00",
        "heure_fin":        "",
        "lieu":             "Centre d'Affaires de Tunis",
        "president_seance": "",
        "redacteur":        "Direction des Risques",
        "introduction":     introduction,
        "participants":     participants,
        "excuses":          [],
        "ordre_du_jour":    odj,
        "points":           points,
        "decisions":        [],
        "plan_action":      plan_action,
        "prochaine_reunion": {
            "date":  prochaine_date,
            "lieu":  "",
            "points_previsionnels": [],
        },
        "approbation": {
            "date":      "",
            "signataire": "",
            "statut":    "En attente",
        },
    }

TARGET_PV_STRUCTURE = {
    "introduction": "",
    
    "ordre_du_jour": [
        {
            "numero": "1",
            "titre": ""
        }
    ],
    
    "points": [
        {
            "numero": "1",
            "titre": "",
            "type": "constat | deliberation | decision",
            "expose": "",
            "discussion": "",
            "conclusion": ""
        }
    ]
}

def export_pv_to_docx(pv: dict, output_path: Optional[str] = None, language: str = "fr") -> str:
    """Exporter un PV structuré en fichier Word (.docx)."""
    if not isinstance(pv, dict):
        raise ValueError("PV doit être un dictionnaire JSON structuré.")

    if "pv" in pv and isinstance(pv["pv"], dict):
        pv = pv["pv"]

    if output_path is None:
        title = pv.get("titre") or pv.get("numero") or "PV"
        safe = re.sub(r"[^\w\-]", "_", str(title))
        output_path = str(EXPORT_DIR / f"{safe}.docx")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin = section.bottom_margin = Inches(0.9)

    def h(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    def para(text, size=11, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(text or "")
        r.bold = bold
        r.font.size = Pt(size)
        return p

    h("PROCÈS-VERBAL DE RÉUNION", size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    if pv.get("titre"):
        h(pv["titre"], size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    metadata = []
    for label, key in [("Date", "date"), ("Heure début", "heure_debut"), ("Heure fin", "heure_fin"), ("Lieu", "lieu"), ("Président", "president_seance"), ("Rédacteur", "redacteur")]:
        if pv.get(key):
            metadata.append(f"{label} : {pv.get(key)}")
    for line in metadata:
        para(line, size=11)
    doc.add_paragraph()

    if pv.get("points"):
        h("POINTS DU PV", size=14)
        for point in pv["points"]:
            h(f"Point {point.get('numero','?')} – {point.get('titre','')}", size=12, bold=True)
            para(point.get("texte_pv", ""), size=11)
            if point.get("type"):
                para(f"Type : {point.get('type')}", size=10, bold=True)
            doc.add_paragraph()

    if pv.get("decisions"):
        h("DÉCISIONS", size=14)
        for decision in pv["decisions"]:
            para(decision, size=11)
        doc.add_paragraph()

    if pv.get("plan_action"):
        h("PLAN D'ACTION", size=14)
        for action in pv["plan_action"]:
            para(f"{action.get('id','')} – {action.get('action','')}", size=11, bold=True)
            details = []
            if action.get("responsable"): details.append(f"Responsable : {action.get('responsable')}")
            if action.get("echeance"): details.append(f"Échéance : {action.get('echeance')}")
            if action.get("statut"): details.append(f"Statut : {action.get('statut')}")
            if details:
                para("; ".join(details), size=10)
            doc.add_paragraph()

    doc.save(output_path)
    return output_path


# ─── Appels LLM ───────────────────────────────────────────────────────────────

SYSTEM_PV = """Tu es un rédacteur administratif expert en procès-verbaux de réunion officiels.

TON RÔLE : transformer des slides de présentation en paragraphes de PV rédigés.

RÈGLE DE TRANSFORMATION PAR TYPE DE SLIDE :

1. SLIDE DE CONSTAT / BILAN → rédige un CONSTAT
   Formules d'ouverture : 
   "La séance débute par l'analyse de...", 
   "Le [titre] présente le bilan de la période...",
   "La Direction prend acte de..."
   → Habille les chiffres avec leur contexte et leur signification.

2. SLIDE DE PLAN / MESURES → rédige une DÉLIBÉRATION  
   Formules d'ouverture :
   "Pour pallier [problème identifié], [le responsable / le Comité] propose...",
   "Il est proposé de mettre en œuvre les mesures suivantes...",
   "Face à ce constat, la Direction envisage..."
   → Chaque mesure devient une proposition argumentée (pas juste une liste).

3. SLIDE DE BUDGET / VOTE → rédige une DÉCISION
   Formules d'ouverture :
   "Après examen et échanges, [le budget / la mesure] est validé(e)...",
   "La Direction approuve à l'unanimité...",
   "Il est décidé de..."
   → Inclure le montant, la date, le résultat du vote.

STYLE OBLIGATOIRE :
- 3ème personne du singulier ou pluriel ("La Direction note", "Le responsable propose")
- Temps : présent de narration ou passé composé
- Jamais de bullet points dans la sortie — uniquement des paragraphes rédigés
- Ton neutre, factuel, administratif
- Chaque slide = 3 à 6 phrases minimum
- Connecteurs logiques entre les idées : "À cet égard,", "En conséquence,", "Il convient de noter que,"

STRUCTURE DE RÉPONSE (JSON) :
{
  "points": [
    {
      "numero": "1",
      "titre": "...",
      "type": "constat | deliberation | decision",
      "texte_pv": "Paragraphe rédigé complet..."
    }
  ],
  "decisions": ["Décision 1 rédigée...", "..."],
  "plan_action": [
    {
      "id": "A01",
      "action": "...",
      "responsable": "...",
      "echeance": "...",
      "statut": "..."
    }
  ]
}
"""


def _call_llm(system: str, user: str, max_tokens: int = 4000) -> str:
    """Appel FORCÉ vers Ollama uniquement."""

    forced_user = (
        user
        + "\n\n⚠️ IMPORTANT: Réponds UNIQUEMENT avec le JSON brut. "
        "Commence par { et termine par }. Aucun autre texte."
    )

    payload = {
        "model": OLLAMA_MODEL,
        "prompt": f"[INST] <<SYS>>\n{system}\n<</SYS>>\n\n{forced_user} [/INST]",
        "stream": False,
        "options": {
            "temperature": 0.1,
            "num_predict": max_tokens,
            "stop": ["\n\n\n"]
        }
    }

    try:
        print("🚀 Appel Ollama...")
        resp = requests.post(
            f"{OLLAMA_URL}/api/generate",
            json=payload        )

        resp.raise_for_status()

        result = resp.json().get("response", "").strip()

        if not result:
            raise ValueError("Réponse vide Ollama")

        return result

    except Exception as e:
        print(f"❌ Erreur Ollama: {e}")
        raise RuntimeError("Ollama indisponible")
# ─── Fonctions principales ────────────────────────────────────────────────────
SYSTEM_PV = """Tu es un expert en gouvernance bancaire et en rédaction de procès-verbaux de comités de direction.
Tu rédiges des PV en français administratif de haut niveau, avec un registre analytique et financier rigoureux.

LEXIQUE OBLIGATOIRE À UTILISER :
- Financier : ratio de solvabilité, covenant, EBITDA, RAROC, NIM (Net Interest Margin), 
  coût du risque, provisions pour créances douteuses, NPL ratio, LCR, NSFR, tier 1, 
  ROE, ROA, cost-to-income, spread de taux, duration, VAR, back-testing, stress test,
  encours de crédit, pipeline commercial, taux de transformation, effet de levier
- Gouvernance : délibération, résolution, quorum, approbation collégiale, mandat,
  pouvoir de décision, habilitation, reporting, escalade, dispositif de contrôle interne
- Analytique : décomposition, ventilation, écart favorable/défavorable, variance,
  glissement annuel, variation trimestrielle, contribution nette, effet prix/volume/mix,
  benchmark sectoriel, trajectoire cible, indicateur avancé

STYLE IMPOSÉ :
- 3ème personne du pluriel impersonnelle ("Le Comité a examiné", "Il a été relevé que")
- Connecteurs analytiques : "L'analyse fait ressortir que", "Il ressort de l'examen que",
  "À l'aune des indicateurs présentés", "Au regard des données probantes",
  "La revue analytique indique", "Les éléments de benchmark confirment"
- Chiffres avec unités précises (M€, pb, %, bps)
- Formulations décisionnelles : "Le Comité a arrêté", "Il a été résolu que",
  "La délibération acte", "Une résolution a été adoptée à l'unanimité"
"""

def _collect_slides_for_odj(odj_titre: str, slides: list) -> list:
    """Retourne toutes les slides rattachées à un item ODJ."""
    target = (odj_titre or "").strip().lower()
    if not target:
        return []
    matched = []
    for slide in slides:
        odj = (slide.get("ordre du jour") or "").strip().lower()
        if odj == target:
            matched.append(slide)
    return matched


def _build_point_data(odj_index: int, odj_titre: str, slides: list) -> str:
    """
    Extrait et formate toutes les données extraites des slides
    qui correspondent à ce point de l'ordre du jour.
    Retourne un bloc texte structuré prêt pour le prompt.
    """
    linked_slides = _collect_slides_for_odj(odj_titre, slides)
    if not linked_slides:
        return f"Aucune donnée extraite pour le point {odj_index}."

    lines = [f"=== Données extraites — Point ODJ {odj_index} : {odj_titre} ==="]

    for slide in linked_slides:
        lines.append(f"\n--- Slide {slide.get('index', '?')} : {slide.get('titre') or 'Sans titre'} ---")

        contenu = [
            c for c in slide.get("contenu", [])
            if isinstance(c, str) and len(c.strip()) > 2 and not c.strip().upper().startswith("POINT")
        ]
        if contenu:
            lines.append("Éléments textuels : " + " | ".join(contenu))

        for g in slide.get("graphiques", []):
            lines.append(f"\nGraphique : {g.get('titre', g.get('type', ''))}")
            lines.append(g.get("resume_pv", ""))

        for img in slide.get("images", []):
            if img:
                lines.append(f"\nImage analysée :\n{img}")

        for t in slide.get("tableaux", []):
            lignes = t.get("lignes", [])
            if not lignes:
                continue
            lines.append(f"\nTableau ({t['nb_lignes']} lignes × {t['nb_colonnes']} colonnes) :")
            header = " | ".join(lignes[0])
            lines.append(header)
            lines.append("-" * len(header))
            for row in lignes[1:]:
                lines.append(" | ".join(row))

    return "\n".join(lines)


def _extract_odj(slides: list) -> list:
    """
    Extrait la liste propre de l'ordre du jour depuis les slides.
    Cherche la slide dont le titre contient 'ordre du jour'.
    """
    for slide in slides:
        titre = (slide.get("titre") or "").lower()
        contenu_join = " ".join(slide.get("contenu", [])).lower()

        if "ordre du jour" in titre or "ordre du jour" in contenu_join:
            raw = slide.get("contenu", [])
            # Filtre les éléments parasites (durées "10 min", numéros seuls)
            propre = [
                c for c in raw
                if len(c) > 8
                and not c.strip().isdigit()
                and not c.strip().lower() in ("ordre du jour",)
                and not c.strip().endswith("min")
                and not c.strip().endswith("min.")
            ]
            return propre
    return []


def generate_pv_draft(client, extracted: dict) -> dict:
    """
    Génère un PV draft structuré à partir du contenu extrait.
    Mappe explicitement chaque point ODJ aux données de sa slide.
    """
    slides = extracted.get("slides", [])
    odj    = _extract_odj(slides)

    if not odj:
        odj = extracted.get("ordre_du_jour", [])

    # Matrice de preuves slide par slide (injectée telle quelle au LLM)
    evidence_matrix = []
    for s in slides:
        evidence_matrix.append({
            "index": s.get("index"),
            "titre": s.get("titre"),
            "ordre_du_jour": s.get("ordre du jour"),
            "contenu": s.get("contenu", []),
            "tableaux": s.get("tableaux", []),
            "graphiques": s.get("graphiques", []),
            "images": s.get("images", []),
            "notes": s.get("notes"),
        })

    # Construit le bloc de données par point ODJ
    points_data_blocks = []
    for i, titre in enumerate(odj):
        bloc = _build_point_data(i + 1, titre, slides)
        points_data_blocks.append(bloc)

    points_data_str = "\n\n".join(points_data_blocks)

    # Participants
    participants_raw = extracted.get("participants", [])
    participants_str = "\n".join(
        f"- {p}" if isinstance(p, str) else f"- {p.get('nom', '')} : {p.get('fonction', '')}"
        for p in participants_raw
    ) or "- [À compléter]"

    # Titre
    titre_reunion = ""
    for s in slides:
        t = s.get("titre") or ""
        if t and not t.isdigit() and len(t) > 5 and "ordre" not in t.lower():
            titre_reunion = t
            break
    if not titre_reunion:
        titre_reunion = extracted.get("meta", {}).get("titre_reunion", "Comité des Risques")

    prompt = f"""
Tu dois générer un procès-verbal COMPLET en JSON strict.
Date du jour : {today}
Titre de la réunion : {titre_reunion}

RÈGLE ABSOLUE ANTI-HALLUCINATION :
- Interdiction d'inventer toute information.
- Utilise EXCLUSIVEMENT les données présentes dans PREUVES STRUCTURÉES.
- Si une information n'est pas disponible (participants, heure, lieu, décisions...), laisse le champ vide ("") ou [].
- N'ajoute jamais de participant, point ODJ, chiffre, action ou décision non présent dans les preuves.

═══════════════════════════════════════
ORDRE DU JOUR ({len(odj)} points) :
═══════════════════════════════════════
{chr(10).join(f"{i+1}. {t}" for i, t in enumerate(odj))}

═══════════════════════════════════════
PARTICIPANTS :
═══════════════════════════════════════
{participants_str}

═══════════════════════════════════════
DONNÉES EXTRAITES PAR POINT ODJ :
═══════════════════════════════════════
{points_data_str}

═══════════════════════════════════════
PREUVES STRUCTURÉES (JSON BRUT SLIDE PAR SLIDE) :
═══════════════════════════════════════
{json.dumps(evidence_matrix, ensure_ascii=False, indent=2)}

═══════════════════════════════════════
RÈGLES DE RÉDACTION OBLIGATOIRES :
═══════════════════════════════════════

1. INTRODUCTION :
   - Ne mentionner QUE des éléments présents dans les preuves.
   - Si date/heure/lieu/participants manquent, ne pas les inventer.

2. POINTS : générer EXACTEMENT {len(odj)} points, un par élément de l'ordre du jour.
   - Chaque point doit utiliser uniquement les slides portant le même "ordre du jour".
   - Si aucune donnée pour un point, écrire explicitement qu'aucune donnée n'est disponible.
   - Chaque point : exposé + discussion + conclusion, basés sur les preuves.
   - Utiliser UNIQUEMENT "Le Comité" comme sujet.
   - Interdiction de bullet points dans les champs texte.
   - Incorporer uniquement les valeurs numériques présentes.
   - Connecteurs obligatoires : "Toutefois", "Par ailleurs", "Ainsi", "En revanche".

3. PLAN D'ACTION :
   - Extraire uniquement les actions explicitement présentes dans les tableaux.
   - Si aucune action explicite, retourner [].

4. FORMAT DE SORTIE = un fichier .docx a telecharger
"""

    raw = _call_llm(SYSTEM_PV, prompt, max_tokens=6000)
    print(f"📨 Réponse LLM (200 chars): {raw[:200]}")

    parsed = _extract_json(raw)
    if parsed:
        # Post-traitement : s'assure que le nombre de points = nombre ODJ
        parsed = _ensure_points_match_odj(parsed, odj, slides)
        print("✅ JSON extrait avec succès")
        return parsed

    print("⚠️ JSON non extrait — fallback structurel")
    return _build_fallback_pv(extracted)


def _ensure_points_match_odj(pv: dict, odj: list, slides: list) -> dict:
    """
    Garantit que le PV a exactement autant de points que l'ODJ.
    Corrige silencieusement si le LLM en a généré plus ou moins.
    """
    points_actuels = pv.get("points", [])

    if len(points_actuels) == len(odj):
        return pv  # parfait

    # Recrée la liste complète
    points_corriges = []
    for i, titre_odj in enumerate(odj):
        if i < len(points_actuels):
            pt = points_actuels[i]
            pt["numero"] = str(i + 1)
            pt["titre"]  = titre_odj
            points_corriges.append(pt)
        else:
            # Point manquant → génère depuis les données brutes
            data = _build_point_data(i + 1, titre_odj, slides)
            points_corriges.append({
                "numero":     str(i + 1),
                "titre":      titre_odj,
                "expose":     f"Dans le cadre du point relatif à {titre_odj}, le Comité a procédé à l'examen des éléments présentés.",
                "discussion": data,
                "conclusion": f"Le Comité a pris acte des informations présentées au titre du point {i+1}.",
            })

    pv["points"] = points_corriges
    return pv


def merge_notes_with_pv(client, pv_draft: dict, notes: list) -> dict:
    """
    Fusion incrémentale des notes réunion dans le PV draft.

    Objectif :
    - conserver le draft existant ;
    - enrichir chaque point ODJ ;
    - reformuler les remarques des participants ;
    - ne jamais réécrire totalement le PV.
    """

    if not notes:
        return pv_draft

    points = pv_draft.get("points", [])

    if not points:
        return pv_draft

    # ─────────────────────────────────────────────
    # Regroupement des notes par ordre du jour
    # ─────────────────────────────────────────────
    notes_by_odj = {}

    for note in notes:

        odj = (
            note.get("ordre_du_jour")
            or note.get("agenda")
            or note.get("point")
            or ""
        ).strip()

        participant = note.get("participant", "").strip()
        content = note.get("content", "").strip()

        if not odj or not content:
            continue

        if odj not in notes_by_odj:
            notes_by_odj[odj] = []

        notes_by_odj[odj].append({
            "participant": participant,
            "content": content
        })

    # ─────────────────────────────────────────────
    # Fusion point par point
    # ─────────────────────────────────────────────
    for point in points:

        titre = (point.get("titre") or "").strip()

        if not titre:
            continue

        linked_notes = notes_by_odj.get(titre)

        if not linked_notes:
            continue

        notes_text = "\n".join(
            f"- {n['participant']} : {n['content']}"
            for n in linked_notes
        )

        existing_discussion = point.get("discussion", "")
        existing_expose = point.get("expose", "")
        existing_conclusion = point.get("conclusion", "")

        prompt = f"""
Tu es un secrétaire de comité bancaire.

TÂCHE UNIQUE : Reformuler les notes suivantes en style administratif formel.
Retourne UNIQUEMENT un JSON avec les remarques reformulées.
Ne touche pas à la discussion existante.
Ne réécris rien d'autre.

NOTES À REFORMULER :
{notes_text}

FORMAT JSON OBLIGATOIRE (rien d'autre) :
{{
  "remarques": [
    "M./Mme [participant] a indiqué que ...",
    "..."
  ]
}}
"""

        try:

            raw = _call_llm(SYSTEM_PV, prompt, max_tokens=2000)

            parsed = _extract_json(raw)

            if not parsed:
                continue

            # ─────────────────────────────────────
            # Mise à jour incrémentale
            # ─────────────────────────────────────

            if parsed.get("discussion"):
                point["discussion"] = parsed["discussion"]

            if parsed.get("conclusion"):
                point["conclusion"] = parsed["conclusion"]

            existing_remarques = point.get("remarques", [])

            if parsed.get("remarques"):
                point["remarques"] = (
                    existing_remarques + parsed["remarques"]
                )

            # décisions globales
            if parsed.get("decisions"):

                if "decisions" not in pv_draft:
                    pv_draft["decisions"] = []

                pv_draft["decisions"].extend(parsed["decisions"])

            # actions globales
            if parsed.get("actions"):

                if "plan_action" not in pv_draft:
                    pv_draft["plan_action"] = []

                for idx, action in enumerate(parsed["actions"], start=1):

                    pv_draft["plan_action"].append({
                        "id": f"A{len(pv_draft['plan_action']) + 1:02d}",
                        "action": action.get("action", ""),
                        "responsable": action.get("responsable", ""),
                        "echeance": action.get("echeance", ""),
                        "statut": "En cours",
                        "priorite": "Normale"
                    })

        except Exception as e:
            print(f"⚠️ Erreur fusion notes pour point '{titre}': {e}")

    return pv_draft

def translate_pv(client, pv: dict, target_language: str) -> dict:
    """Traduire le PV dans la langue cible."""

    lang_names = {
        "fr": "français",
        "en": "English (formal administrative register)",
        "ar": "arabe (registre administratif officiel, écriture RTL)"
    }
    lang_name = lang_names.get(target_language, target_language)

    prompt = f"""Traduis ce PV en {lang_name}.

PV:
{json.dumps(pv, ensure_ascii=False, indent=2)}

RÈGLES:
- Traduis tous les champs textuels (titre, expose, discussion, remarques, conclusion, decisions, actions...)
- Conserve les dates, numéros, codes (PV-2025-001, A01...) tels quels
- Maintiens le registre formel et administratif
- Conserve EXACTEMENT la même structure JSON
- RÉPONSE = JSON UNIQUEMENT, commence par {{"""

    raw = _call_llm(SYSTEM_PV, prompt, max_tokens=4000)
    parsed = _extract_json(raw)

    if parsed:
        parsed["_language"] = target_language
        return parsed

    print("⚠️ Traduction JSON non extraite — retour de l'original")
    return pv