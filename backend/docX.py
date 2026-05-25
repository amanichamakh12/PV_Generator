from generate_pv_draft import generate_pv_draft_pipeline
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Any

def build_pv_request_body(
    extracted: dict[str, Any],
    use_llm_for_slides: bool = False,
    use_llm_for_analysis: bool = False,
) -> dict[str, Any]:
    """
    Transforme la sortie de parse_pptx en body prêt pour /api/test-draft-pipeline.
    Fait tourner generate_pv_draft_pipeline en interne et mappe les clés
    vers le format attendu par build_pv_docx.
    """

    # ── 1. Pipeline brut ───────────────────────────────────────────────────
    raw = generate_pv_draft_pipeline(
        extracted,
        use_llm_for_slides=use_llm_for_slides,
        use_llm_for_analysis=use_llm_for_analysis,
    )

    # ── 2. Metadata depuis les slides de couverture ────────────────────────
    metadata = _extract_metadata(extracted.get("slides") or [])

    # ── 3. Mapping analyses : pipeline → build_pv_docx ────────────────────
    # Pipeline produit  : analyse, constats, risques, actions_suggerees
    # build_pv_docx attend : resume, points_cles, decision
    analyses_mappees = []
    for item in raw.get("analyses_par_ordre_du_jour") or []:
        odj = item.get("ordre_du_jour") or "—"

        # Saute les slides sans ordre du jour réel (couverture, fin, etc.)
        if odj in ("hors ordre du jour",):
            continue

        # resume  ← analyse consolidée
        resume = item.get("analyse") or ""

        # points_cles ← constats + risques (dédoublonnés)
        points_cles = _dedupe(
            (item.get("constats") or []) +
            (item.get("risques") or [])
        )

        # decision ← première action suggérée non vide
        actions = [a for a in (item.get("actions_suggerees") or []) if a.strip()]
        decision = actions[0] if actions else ""

        analyses_mappees.append({
            "ordre_du_jour": odj,
            "resume":        resume,
            "points_cles":   points_cles,
            "decision":      decision,
        })

    # ── 4. Mapping groupes : on garde la structure existante ───────────────
    groupes_mappes = []
    for g in raw.get("groupes_ordre_du_jour") or []:
        odj = g.get("ordre_du_jour") or "—"
        if odj == "hors ordre du jour":
            continue
        groupes_mappes.append({
            "ordre_du_jour": odj,
            "paragraphes":   g.get("paragraphes") or [],
        })

    # ── 5. Body final ──────────────────────────────────────────────────────
    return {
        "extracted": {                         # clé attendue par DraftPipelineRequest
            "metadata":                    metadata,
            "analyses_par_ordre_du_jour":  analyses_mappees,
            "groupes_ordre_du_jour":       groupes_mappes,
        },
        "use_llm_for_slides":   use_llm_for_slides,
        "use_llm_for_analysis": use_llm_for_analysis,
    }


# ── Helpers privés ─────────────────────────────────────────────────────────────

def _dedupe(items: list[str]) -> list[str]:
    seen, out = set(), []
    for item in items:
        key = item.strip().lower()
        if key and key not in seen:
            seen.add(key)
            out.append(item.strip())
    return out


def _extract_metadata(slides: list[dict]) -> dict:
    """
    Lit les slides 1 (couverture) et 2 (ordre du jour)
    pour reconstituer le bloc metadata sans rien coder en dur.
    """
    meta = {
        "titre":       "",
        "date":        "",
        "heure":       "",
        "lieu":        "",
        "institution": "",
        "comite":      "",
        "president":   "son Président",
        "presents":    [],
    }

    for slide in slides:
        idx     = slide.get("index")
        contenu = slide.get("contenu") or []
        titre   = (slide.get("titre") or "").strip()

        if idx == 1:
            meta["comite"] = titre  # ex: "COMITÉ DES RISQUES"

            for line in contenu:
                line = line.strip()
                low  = line.lower()

                # "Tableau de Bord Risques — 1er Trimestre 2025"
                if "tableau de bord" in low or "trimestre" in low or "semestre" in low:
                    meta["titre"] = f"PV de Comité des risques"

                # "Date : Avril 2025   │   Direction des Risques"
                if "date" in low or "│" in line:
                    for part in line.split("│"):
                        part = part.strip()
                        low_part = part.lower()
                        if low_part.startswith("date"):
                            meta["date"] = part.split(":", 1)[-1].strip()
                        if "direction" in low_part or "banque" in low_part:
                            meta["institution"] = part

    # Fallback titre
    if not meta["titre"] and meta["comite"]:
        meta["titre"] = f"PV de {meta['comite'].title()}"

    return meta
def build_pv_docx(pipeline_result: dict) -> bytes:
    doc = DocxDocument()

    # ── Style de base ────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    meta        = pipeline_result.get("metadata") or {}
    analyses    = pipeline_result.get("analyses_par_ordre_du_jour") or []
    groupes     = pipeline_result.get("groupes_ordre_du_jour") or []

    titre       = meta.get("titre",       "Procès-Verbal de Réunion")
    date        = meta.get("date",        "—")
    heure       = meta.get("heure",       "—")
    lieu        = meta.get("lieu",        "—")
    institution = meta.get("institution", "l'institution")
    comite      = meta.get("comite",      "le comité")
    president   = meta.get("president",   "son Président")
    presents    = meta.get("presents",    [])

    groupes_by_odj = {g["ordre_du_jour"]: g for g in groupes}

    # ── 1. Titre centré + ligne de séparation ────────────────────────────
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titre.add_run(titre)
    run.bold = True
    run.font.size = Pt(14)



    def add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    add_bottom_border(p_titre)
    doc.add_paragraph()  # espace

    # ── 2. Phrase d'introduction ─────────────────────────────────────────
    intro = (
        f"Le {comite} de {institution} s'est réuni le {date} à {heure} au {lieu}, "
        f"sur invitation de {president}, pour délibérer sur l'ordre du jour suivant :"
    )
    doc.add_paragraph(intro)

    # ── 3. Ordre du jour numéroté bold ───────────────────────────────────
    for i, analyse in enumerate(analyses, 1):
        odj = analyse.get("ordre_du_jour", "—")
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(odj)
        run.bold = True

    doc.add_paragraph()  # espace

    # ── 4. Présents ──────────────────────────────────────────────────────
    if presents:
        doc.add_paragraph("Etaient présents :")
        for membre in presents:
            doc.add_paragraph(membre, style="List Bullet")
        doc.add_paragraph()

    # ── 5. Formalités d'ouverture ────────────────────────────────────────
    doc.add_paragraph(
        "La feuille de présence a été établie et dûment signée par les membres présents."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Apres avoir constaté la présence requise des membres, le président a ouvert "
        "les travaux de la comité en souhaitant la bienvenue aux présents ."
    )
    doc.add_paragraph()

    # ── 6. Sections par point de l'ordre du jour ─────────────────────────
    for i, analyse in enumerate(analyses, 1):
        odj       = analyse.get("ordre_du_jour", "—")
        resume    = analyse.get("resume", "")
        points    = analyse.get("points_cles") or []
        decision  = analyse.get("decision", "")
        groupe    = groupes_by_odj.get(odj, {})
        paragraphes = groupe.get("paragraphes") or []

        # Titre de section bold
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {odj}")
        run.bold = True

        # Résumé
        if resume:
            doc.add_paragraph(resume)

        # Paragraphes extraits des slides
        for para in paragraphes:
            if para:
                doc.add_paragraph(para)

        # Points clés en liste à puces
        for pt in points:
            doc.add_paragraph(pt, style="List Bullet")

        # Décision
        if decision:
            p = doc.add_paragraph()
            p.add_run("Décision : ").bold = True
            p.add_run(decision)

        doc.add_paragraph()  # espace entre sections

    # ── Sérialisation ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
